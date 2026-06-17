# src/nodes/assembly_node.py
import os
import re
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, Dict, Any
from src.state import GraphState

# =====================================================
# Helper Functions
# =====================================================
def get_client_name(questionnaire):
    if isinstance(questionnaire, list) and len(questionnaire) > 0:
        questionnaire = questionnaire[0]
    if isinstance(questionnaire, dict):
        return (questionnaire.get("company_name") or
                questionnaire.get("client_name") or
                questionnaire.get("organization") or
                questionnaire.get("company") or
                "Client")
    return "Client"

def generate_proposal_filename(state: GraphState) -> str:
    questionnaire = state.get("questionnaire", {})
    client_name = get_client_name(questionnaire)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = "".join(c for c in client_name if c.isalnum() or c in " ._-").strip()
    clean_name = clean_name.replace(" ", "_")
    return f"Proposal_{clean_name}_{timestamp}"

def generate_proposal_summary(state: GraphState) -> Dict[str, Any]:
    sections_completed = state.get("sections_completed", [])
    total_chunks = 0
    section_metrics = {}
    section_order = [
        "business_context", "overview", "understanding", "objectives",
        "deliverables", "approach", "outcomes", "business_impact"
    ]
    for section_key in section_order:
        section_data = state.get(section_key)
        if section_data and isinstance(section_data, dict):
            content = section_data.get("content", "")
            section_metrics[section_key] = {
                "chunks_used": section_data.get("chunks_used", 0),
                "has_content": bool(content),
                "content_length": len(content)
            }
            total_chunks += section_data.get("chunks_used", 0)
    questionnaire = state.get("questionnaire", {})
    client_name = get_client_name(questionnaire)
    summary = {
        "client_name": client_name,
        "generated_date": datetime.now().isoformat(),
        "sections_completed": sections_completed,
        "total_sections": len(sections_completed),
        "total_chunks_used": total_chunks,
        "section_metrics": section_metrics,
        "metadata": state.get("metadata_dict", {})
    }
    summary_path = "x_results/proposal_summary.json"
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)
    print(f"📊 Proposal summary saved to: {summary_path}")
    return summary

# =====================================================
# JMAN brand colours
# =====================================================
_NAVY    = RGBColor(0x17, 0x38, 0x45)
_PINK    = RGBColor(0xFF, 0x61, 0x96)
_DARK    = RGBColor(0x1D, 0x1C, 0x1C)
_GRAY    = RGBColor(0x80, 0x80, 0x80)
_DEEP    = RGBColor(0x19, 0x10, 0x5B)
_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
_BORDER  = RGBColor(0xDD, 0xDD, 0xDD)
_ZEBRA   = "F9F9F9"
_RED     = RGBColor(0xCC, 0x00, 0x00)
_AMBER   = RGBColor(0xCC, 0x77, 0x00)
_GREEN   = RGBColor(0x2E, 0x7D, 0x32)

_FONT = "Arial"

# =====================================================
# Brand asset locations
# =====================================================
_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "jman")
_ASSET_FILES = {
    "logo_white":  "jman_logo_white.png",
    "logo_navy":   "jman_logo_navy.png",
    "corner_mark": "jman_corner_mark.png",
    "cover_bg":    "cover_background.jpg",
    "back_cover_bg": "back_cover_background.jpg",
}

def _asset(name: str) -> Optional[str]:
    path = os.path.join(_ASSETS_DIR, _ASSET_FILES[name])
    return path if os.path.isfile(path) else None

# =====================================================
# Page geometry (A4)
# =====================================================
_PAGE_W_CM = 21.0
_PAGE_H_CM = 29.7

# =====================================================
# Low-level helpers
# =====================================================

def _run(para, text, size_pt, bold=False, italic=False, color=None, font=_FONT):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def _set_cell_borders(cell, color_hex="DDDDDD", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _set_cell_margins(cell, top=80, bottom=80, left=110, right=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def _set_table_width(tbl, width_twips: int):
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    tblStyle = tblPr.find(qn("w:tblStyle"))
    if tblStyle is not None:
        tblStyle.addnext(tblW)
    else:
        tblPr.insert(0, tblW)

def _pink_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF6196")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    return p

def _spacer(doc, pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    return p

# =====================================================
# Floating (page-anchored) pictures
# =====================================================
def _add_floating_picture(paragraph, image_path, width, height, behind_doc=True, name="Picture"):
    run = paragraph.add_run()
    run.add_picture(image_path, width=width, height=height)
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))
    inline.remove(extent)
    inline.remove(docPr)
    inline.remove(graphic)
    docPr.set("name", name)

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc", "1" if behind_doc else "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simplePos = OxmlElement("wp:simplePos")
    simplePos.set("x", "0")
    simplePos.set("y", "0")

    posH = OxmlElement("wp:positionH")
    posH.set("relativeFrom", "page")
    offH = OxmlElement("wp:posOffset")
    offH.text = "0"
    posH.append(offH)

    posV = OxmlElement("wp:positionV")
    posV.set("relativeFrom", "page")
    offV = OxmlElement("wp:posOffset")
    offV.text = "0"
    posV.append(offV)

    effectExtent = OxmlElement("wp:effectExtent")
    for side in ("l", "t", "r", "b"):
        effectExtent.set(side, "0")

    wrapNone = OxmlElement("wp:wrapNone")

    anchor.append(simplePos)
    anchor.append(posH)
    anchor.append(posV)
    anchor.append(extent)
    anchor.append(effectExtent)
    anchor.append(wrapNone)
    anchor.append(docPr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)
    return run

# =====================================================
# Styled content parsing and application
# =====================================================
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

def _add_inline_runs(p, text, size_pt, color, bold=False, italic=False):
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _run(p, text[pos:m.start()], size_pt, bold=bold, italic=italic, color=color)
        _run(p, m.group(1), size_pt, bold=True, italic=italic, color=color)
        pos = m.end()
    if pos < len(text):
        _run(p, text[pos:], size_pt, bold=bold, italic=italic, color=color)

def _status_color(value: str):
    v = value.strip().lower()
    if any(k in v for k in ("blocked", "at risk", "failed", "delayed")):
        return _RED
    if any(k in v for k in ("in progress", "pending", "ongoing", "review")):
        return _AMBER
    if any(k in v for k in ("resolved", "done", "complete", "approved", "on track")):
        return _GREEN
    return _DARK

def _render_table(doc, header_cells, body_rows):
    n_cols = len(header_cells)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    content_width_twips = int((_PAGE_W_CM - 3.5) * 567)
    if n_cols == 2:
        widths = [int(content_width_twips * 0.7), int(content_width_twips * 0.3)]
    else:
        widths = [content_width_twips // n_cols] * n_cols
    _set_table_width(tbl, sum(widths))

    hdr_row = tbl.rows[0]
    for i, text in enumerate(header_cells):
        cell = hdr_row.cells[i]
        cell.width = Pt(widths[i] / 20)
        _set_cell_shading(cell, "173845")
        _set_cell_borders(cell, "DDDDDD", 4)
        _set_cell_margins(cell)
        cell.paragraphs[0].clear() if cell.paragraphs[0].runs else None
        p = cell.paragraphs[0]
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        _run(p, text, 10, bold=True, color=_WHITE)

    for ridx, row_vals in enumerate(body_rows):
        row = tbl.add_row()
        fill = "FFFFFF" if ridx % 2 == 0 else _ZEBRA
        for i, val in enumerate(row_vals):
            cell = row.cells[i]
            cell.width = Pt(widths[i] / 20)
            _set_cell_shading(cell, fill)
            _set_cell_borders(cell, "DDDDDD", 4)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            is_status_col = (n_cols == 2 and i == 1)
            color = _status_color(val) if is_status_col else _DARK
            _run(p, val, 10, bold=is_status_col, color=color)
    return tbl

def _apply_styled_content(doc, raw_text: str, heading_counter: list):
    lines = raw_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|?[\s:|-]+\|?$", lines[i + 1].strip()):
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            j = i + 2
            body_rows = []
            while j < n and lines[j].strip().startswith("|"):
                body_rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            _render_table(doc, header_cells, body_rows)
            _spacer(doc, pt=8)
            i = j
            continue

        if stripped.startswith("# "):
            heading_counter[0] += 1
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.02)
            pf.first_line_indent = Cm(-1.02)
            pf.space_before = Pt(14)
            pf.space_after = Pt(8)
            try:
                pf.tab_stops.add_tab_stop(Cm(1.02))
            except Exception:
                pass
            _run(p, f"{heading_counter[0]}\t", 12, bold=True, color=_DEEP)
            _add_inline_runs(p, text, 12, _DEEP, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            _add_inline_runs(p, text, 11, _PINK, bold=True)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.27)
            pf.first_line_indent = Cm(-0.64)
            pf.space_after = Pt(4)
            pf.line_spacing = 1.15
            try:
                pf.tab_stops.add_tab_stop(Cm(1.27))
            except Exception:
                pass
            _run(p, "•\t", 10, color=_PINK)
            _add_inline_runs(p, text, 10, _DARK)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        _add_inline_runs(p, stripped, 10, _DARK)
        i += 1

# =====================================================
# Cover page
# =====================================================
def _build_cover_section(doc, client_name: str):
    sec = doc.sections[0]
    sec.page_width = Cm(_PAGE_W_CM)
    sec.page_height = Cm(_PAGE_H_CM)
    sec.top_margin = Cm(1.0)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.75)
    sec.right_margin = Cm(1.75)

    bg = _asset("cover_bg")
    if bg:
        p_bg = doc.add_paragraph()
        p_bg.paragraph_format.space_after = Pt(0)
        _add_floating_picture(
            p_bg, bg, Cm(_PAGE_W_CM), Cm(_PAGE_H_CM), behind_doc=True, name="Cover Background"
        )

    logo = _asset("logo_white")
    if logo:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_logo.paragraph_format.space_after = Pt(0)
        r = p_logo.add_run()
        r.add_picture(logo, width=Cm(8.0))

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(210)

    p_rule = doc.add_paragraph()
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "FFFFFF")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p_rule.paragraph_format.space_after = Pt(12)

    p_caption = doc.add_paragraph()
    p_caption.paragraph_format.space_after = Pt(14)
    _run(p_caption, "STATEMENT OF WORK BETWEEN", 15, bold=True, color=_PINK)

    p_company = doc.add_paragraph()
    p_company.paragraph_format.space_after = Pt(6)
    _run(p_company, client_name, 30, bold=True, color=_NAVY)

    p_and = doc.add_paragraph()
    p_and.paragraph_format.space_after = Pt(6)
    _run(p_and, "And", 16, italic=True, color=_GRAY)

    p_jman = doc.add_paragraph()
    p_jman.paragraph_format.space_after = Pt(0)
    _run(p_jman, "JMAN Group", 30, bold=True, color=_NAVY)

# =====================================================
# Back cover (no longer used – kept only for reference)
# =====================================================
# def _build_back_cover_section(doc):
#     sec = doc.add_section(WD_SECTION.NEW_PAGE)
#     ...

# =====================================================
# Helper to insert a PAGE field
# =====================================================
def _add_page_number_field(paragraph, size_pt=8, color=_GRAY):
    """Insert a PAGE field into the given paragraph (new run)."""
    run = paragraph.add_run()
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    fld = OxmlElement('w:fld')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fld.append(instr)
    run._r.append(fld)

# =====================================================
# Content section: header (logo), footer (client, version, date, page, mark)
# =====================================================
def _build_content_section(doc, client_name="Client", version="v1.0"):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(_PAGE_W_CM)
    sec.page_height = Cm(_PAGE_H_CM)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(0.9)
    sec.left_margin = Cm(1.75)
    sec.right_margin = Cm(1.75)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(0.9)
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False

    # Header
    header = sec.header
    for p in list(header.paragraphs):
        p.clear()
    p_logo = header.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_logo.paragraph_format.space_after = Pt(0)
    logo = _asset("logo_navy")
    if logo:
        r = p_logo.add_run()
        r.add_picture(logo, width=Cm(3.2))

    # Footer
    footer = sec.footer
    for p in list(footer.paragraphs):
        p.clear()

    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(4)
    fp.paragraph_format.space_after = Pt(0)

    # Navy top border
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_bdr = OxmlElement("w:top")
    top_bdr.set(qn("w:val"), "single")
    top_bdr.set(qn("w:sz"), "8")
    top_bdr.set(qn("w:space"), "4")
    top_bdr.set(qn("w:color"), "173845")
    pBdr.append(top_bdr)
    pPr.append(pBdr)

    # Tab stops
    fp.paragraph_format.tab_stops.clear_all()
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(0), WD_TAB_ALIGNMENT.LEFT)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(17.5), WD_TAB_ALIGNMENT.RIGHT)

    month_year = datetime.now().strftime("%B %Y")

    # Left: client
    _run(fp, f"Circulation Limited: {client_name}", 8, color=_GRAY)
    _run(fp, "\t", 8)

    # Centre: version
    _run(fp, f"Version {version}", 8, color=_GRAY)
    _run(fp, "\t", 8)

    # Right: month‑year, page number, corner mark
    _run(fp, month_year, 8, italic=True, color=_GRAY)
    _run(fp, " ", 8)

    # Page number field
    _add_page_number_field(fp, 8, _GRAY)

    _run(fp, " ", 8)

    # Corner mark
    mark = _asset("corner_mark")
    if mark:
        r = fp.add_run()
        r.add_picture(mark, height=Cm(0.45))

    return sec

# =====================================================
# Main LangGraph Node
# =====================================================
def assemble_proposal_node(state: GraphState) -> GraphState:
    print("\n" + "=" * 80)
    print("📄 ASSEMBLING: Complete Proposal")
    print("=" * 80)

    questionnaire = state.get("questionnaire", {})
    client_name = get_client_name(questionnaire)
    version = state.get("version", "v1.0")

    try:
        doc = Document()

        # 1. Cover page
        _build_cover_section(doc, client_name)

        # 2. Content section (header + footer + body)
        _build_content_section(doc, client_name, version)

        # 3. Body content
        section_order = [
            ("business_context", "Business Context"),
            ("overview",         "Overview"),
            ("understanding",    "Understanding"),
            ("objectives",       "Objectives"),
            ("deliverables",     "Deliverables"),
            ("approach",         "Approach"),
            ("outcomes",         "Outcomes"),
            ("business_impact",  "Business Impact"),
        ]

        heading_counter = [0]
        for section_key, section_title in section_order:
            print(f"\n📄 Processing: {section_title}")
            section_data = state.get(section_key)
            if not section_data or not isinstance(section_data, dict):
                print(f"   ⚠️ No data for {section_title}, skipping...")
                continue
            content = section_data.get("content", "")
            if not content:
                print(f"   ⚠️ Empty content for {section_title}, skipping...")
                continue
            print(f"   ✅ Adding content: {len(content)} characters")
            _apply_styled_content(doc, content, heading_counter)
            _pink_divider(doc)

        # Back cover removed

        # Save
        filename = generate_proposal_filename(state)
        os.makedirs("x_results", exist_ok=True)
        word_path = f"x_results/{filename}.docx"
        doc.save(word_path)
        print(f"\n✅ Word document saved to: {word_path}")

        summary = generate_proposal_summary(state)

        state["proposal"] = {
            "filename": filename,
            "word_path": word_path,
            "summary": summary,
            "timestamp": datetime.now().isoformat()
        }

        print("\n" + "=" * 80)
        print("✅ PROPOSAL ASSEMBLY COMPLETE!")
        print("=" * 80)
        print(f"📄 Word Doc: {word_path}")
        print(f"📊 Summary: x_results/proposal_summary.json")
        print("=" * 80)

    except Exception as e:
        print(f"❌ Error during assembly: {e}")
        import traceback
        traceback.print_exc()
        state["error"] = f"Assembly failed: {str(e)}"

    state["sections_completed"].append("Proposal Assembly")
    return state